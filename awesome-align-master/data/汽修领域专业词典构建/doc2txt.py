#!/usr/bin/env python
# -*- coding: UTF-8 -*-
"""
@Project ：awesome-align-master 
@File    ：doc2txt.py
@Author  ：znn
@Date    ：2022/9/14 13:28 
"""
import os
import docx
base_dir = os.getcwd()
file_name = "英汉机电工程缩略语词典" #"英汉机电工程缩略语词典"或者“图解英汉汽车技术词典_第3版”
# 每个循环中docx文档和txt文档的命名
files_dir_path = os.path.join(base_dir,file_name,"docx")
# 新建和打开txt文档
for file in os.listdir(files_dir_path):
    if file.split(".")[-1] == "docx":
        filenamedocx = os.path.join(files_dir_path,file)
        filenametxt = os.path.join(files_dir_path,"../txt",str("".join(file.split(".")[:-1])+"txt"))
        f = open(filenametxt, 'w+',encoding="utf-8")
        # 打开docx的文档并读入名为file的变量中
        file = docx.Document(filenamedocx)
        # 输入docx中的段落数，以检查是否空文档
        print('段落数:' + str(len(file.paragraphs)))
        # 将每个段落的内容都写进去txt里面
        for para in file.paragraphs:
            f.write(para.text+"\n")
        f.close()


